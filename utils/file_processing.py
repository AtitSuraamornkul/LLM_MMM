import PyPDF2
import docx
import io
from datetime import datetime
import pandas as pd
import tempfile
import os

import streamlit as st
from langchain_experimental.agents import create_csv_agent
from langchain_ollama import ChatOllama
from langchain_groq import ChatGroq

OLLAMA_BASE_URL = os.getenv("OLLAMA_BASE_URL")
OLLAMA_MODEL = "llama3.1:latest" 

def process_csv_with_agent(uploaded_file, user_prompt):
    """Process CSV with agent based on user prompt - always requires a prompt"""
    try:
        with tempfile.NamedTemporaryFile(mode='w+b', suffix='.csv', delete=False) as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_file_path = tmp_file.name
        
        try:
            llm = ChatOllama(
                  model=OLLAMA_MODEL,  # You can change this to your preferred model
                  temperature=0,
                  base_url=OLLAMA_BASE_URL
                  #base_url="http://localhost:11434",  # Default Ollama URL
                  # timeout=300,  # 5 minutes timeout for complex analysis
              )
            # llm = ChatGroq(
            #      model="llama-3.3-70b-versatile",  # You can change this to your preferred model
            #      temperature=0,
            #      #base_url="http://localhost:11434",  # Default Ollama URL
            #      #timeout=300,  # 5 minutes timeout for complex analysis
            #  )

            agent = create_csv_agent(
                llm,
                tmp_file_path,
                verbose=False,  # Reduced verbosity
                allow_dangerous_code=True,
                handle_parsing_errors=True  
            )

            # Always use the user prompt for analysis
            analysis_prompt = f"""
            The date format is in Year-Month-Date unless specified
            All spendings and revenue is in Thai Baht (THB)

            Based on the user's question: "{user_prompt}"

            PREPROCESSING STEPS:
                1. First, examine the dataframe structure with df.info() and df.head()
                2. Identify numeric columns only for correlation analysis
                3. For date columns, keep them as datetime objects, do NOT convert to int64
                4. Use df.select_dtypes(include=[np.number]) to get only numeric columns for correlation

            Analyze the CSV file ({uploaded_file.name}) to answer this question.
            Provide specific insights, statistics, and findings relevant to the question.
            Include any relevant data patterns, trends, or anomalies you discover.
            If you need to work with dates, use pd.to_datetime() but don't convert to int64 for correlation

            After preprocessing, use the contents of the CSV file ({uploaded_file.name}) to extract and report all relevant facts, statistics, and comparisons required to answer the user’s question.
            If the user’s question asks for a comparison (explicitly or implicitly) between the CSV and another source (e.g., MMM report), extract the required metric(s) from both and present a direct, factual comparison.
            If the question is not directly about the CSV, but the answer can be supported or enriched with CSV data, report the relevant numbers and insights from the CSV as supporting evidence.
            Never stop at data format or preprocessing details—always continue to deliver business insights, trends, or comparisons as appropriate.
            
            if encounter acronyms, don't assume the meaning, refer to it as the given acronym

            DO NOT provide any estimates or predictions only facts
            IGNORE index limit and index column
            """
            
            result = agent.invoke({"input": analysis_prompt})
            
            # Extract output properly
            if isinstance(result, dict):
                output = result.get("output", str(result))
            else:
                output = str(result)

            formatted_result = f"""
=== CSV ANALYSIS: {uploaded_file.name} ===
Question: {user_prompt}

{output}

=== END CSV ANALYSIS ===
"""
                
            return formatted_result

        finally:
            # Clean up temporary file
            if os.path.exists(tmp_file_path):
                os.unlink(tmp_file_path)
                
    except Exception as e:
        return f"Error analyzing CSV {uploaded_file.name}: {str(e)}"


def extract_file_content(uploaded_file, user_prompt=None):
    """Extract content from files - CSV processing requires user_prompt"""
    try:
        if uploaded_file.type == "text/plain":
            content = str(uploaded_file.read(), "utf-8")
            return content
        
        elif uploaded_file.type == "application/pdf":
            # PDF extraction
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = ""
            
            # Process all pages
            for page_num in range(len(pdf_reader.pages)):
                try:
                    page_text = pdf_reader.pages[page_num].extract_text()
                    if page_text.strip():
                        text += f"--- Page {page_num + 1} ---\n{page_text}\n\n"
                except Exception as e:
                    text += f"--- Page {page_num + 1} ---\n[Error extracting page: {str(e)}]\n\n"
            
            return text
        
        elif uploaded_file.type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]:
            # Word document extraction
            doc = docx.Document(uploaded_file)
            text = ""
            
            # Extract all paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract all tables
            if doc.tables:
                text += "\n--- Document Tables ---\n"
                for i, table in enumerate(doc.tables):
                    text += f"Table {i+1}:\n"
                    for row in table.rows:
                        row_text = " | ".join([cell.text.strip() for cell in row.cells])
                        text += f"{row_text}\n"
                    text += "\n"
            
            return text
        
        elif uploaded_file.type == "text/csv":
            if user_prompt:
                return process_csv_with_agent(uploaded_file, user_prompt)
            else:
                return f"CSV file {uploaded_file.name} ready for analysis. Please ask a question to analyze this data."
        
        else:
            return f"Unsupported file type: {uploaded_file.type}\nFile: {uploaded_file.name}\nSupported formats: Text (.txt), PDF (.pdf), Word (.docx, .doc), CSV (.csv)"
            
    except Exception as e:
        return f"Error processing {uploaded_file.name}: {str(e)}"



def analyze_csv_files_dynamically(user_prompt):
    """Analyze all uploaded CSV files with the current user prompt"""
    if not hasattr(st.session_state, 'uploaded_csv_files') or not st.session_state.uploaded_csv_files:
        return ""
    
    csv_analyses = []
    
    for csv_file in st.session_state.uploaded_csv_files:
        try:
            # Reset file pointer to beginning
            csv_file.seek(0)
            
            # Analyze with current prompt
            analysis = process_csv_with_agent(csv_file, user_prompt)
            csv_analyses.append(analysis)
            
        except Exception as e:
            error_analysis = f"=== CSV ANALYSIS ERROR: {csv_file.name} ===\nError: {str(e)}\n=== END ERROR ==="
            csv_analyses.append(error_analysis)
    
    return "\n".join(csv_analyses)
    

def process_uploaded_files(uploaded_files, user_prompt=None):
    """Process files, storing CSV files separately for dynamic analysis"""
    context_parts = []
    csv_files = []
    
    for file in uploaded_files:
        if file.type == "text/csv":
            # Store CSV files for dynamic processing
            csv_files.append(file)
            # Don't add placeholder text to static context
            # context_parts.append(f"=== CSV FILE: {file.name} ===\nCSV file uploaded and ready for dynamic analysis.\n")
        else:
            # Process static files normally
            content = extract_file_content(file, None)
            context_parts.append(f"=== FILE: {file.name} ===\n{content}\n")
    
    # Store CSV files in session state for dynamic processing
    if csv_files:
        st.session_state.uploaded_csv_files = csv_files
    
    # Only return static file context
    return "\n".join(context_parts)